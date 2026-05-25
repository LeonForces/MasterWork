from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def set_run_font(run, size_pt: int, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is not None:
        rfonts.set(qn("w:ascii"), "Times New Roman")
        rfonts.set(qn("w:hAnsi"), "Times New Roman")
        rfonts.set(qn("w:eastAsia"), "Times New Roman")
        rfonts.set(qn("w:cs"), "Times New Roman")


def insert_page_break_after(paragraph) -> None:
    page_break_paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run.append(br)
    page_break_paragraph.append(run)
    paragraph._p.addnext(page_break_paragraph)


def main(path: str) -> None:
    doc_path = Path(path)
    doc = Document(doc_path)

    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    in_title_page = True
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style is not None else ""
        fmt = paragraph.paragraph_format
        text = paragraph.text.strip()
        if in_title_page:
            fmt.first_line_indent = Cm(0)
            fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fmt.line_spacing = 1.0
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(6)
            for run in paragraph.runs:
                set_run_font(run, 14, True if text in {"ОТЧЕТ", "ПО МАГИСТЕРСКОЙ ДИПЛОМНОЙ РАБОТЕ"} else None)
            if text.replace("\u00a0", " ") == "Москва 2026":
                insert_page_break_after(paragraph)
                in_title_page = False
            continue

        if style_name.startswith("Heading 1"):
            fmt.first_line_indent = Cm(0)
            fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
            fmt.line_spacing = 1.0
            fmt.space_before = Pt(18)
            fmt.space_after = Pt(12)
            fmt.page_break_before = True
            for run in paragraph.runs:
                run.text = run.text.upper()
                set_run_font(run, 14, True)
        elif style_name.startswith("Heading"):
            fmt.first_line_indent = Cm(0)
            fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
            fmt.line_spacing = 1.0
            fmt.space_before = Pt(12)
            fmt.space_after = Pt(6)
            for run in paragraph.runs:
                set_run_font(run, 14, True)
        elif "Caption" in style_name or paragraph.text.strip().startswith(("Таблица", "Рисунок")):
            fmt.first_line_indent = Cm(0)
            fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
            fmt.line_spacing = 1.0
            fmt.space_before = Pt(12)
            fmt.space_after = Pt(6)
            for run in paragraph.runs:
                set_run_font(run, 14)
        else:
            fmt.first_line_indent = Cm(1.25)
            fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            fmt.line_spacing = 1.5
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(0)
            for run in paragraph.runs:
                set_run_font(run, 14)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.first_line_indent = Cm(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        set_run_font(run, 12)

    doc.save(doc_path)


if __name__ == "__main__":
    if len(sys.argv) != 2:
        raise SystemExit("Usage: postprocess_docx.py <docx>")
    main(sys.argv[1])
